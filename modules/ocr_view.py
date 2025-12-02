import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF
from PIL import Image
import io
from docx import Document
import re
import pandas as pd

def get_available_models(api_key):
    try:
        genai.configure(api_key=api_key)
        all_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        return all_models
    except:
        return []

def clean_ocr_text(text):
    if not text: return ""
    # ลบ Markdown code block ออก
    text = text.replace("```csv", "").replace("```", "")
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        if re.match(r'^[\s\|\-\_\=\:\+]{5,}$', line.strip()):
            continue
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)

# --- ปรับปรุงให้รองรับ CSV ---
def ocr_single_image(api_key, image, model_name, output_format="text"):
    try:
        genai.configure(api_key=api_key)
        
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
        
        if output_format == "csv":
            # Prompt สำหรับ Excel
            prompt = """
            Act as a Data Entry Clerk. 
            Extract the table data from this image perfectly.
            - Output STRICTLY in CSV format (Comma Separated Values).
            - Do NOT use Markdown code blocks. Just raw CSV data.
            - Handle Thai characters correctly.
            - If there are merged cells, repeat the value in each cell or handle logically.
            """
        else:
            # Prompt สำหรับ Word (Text)
            prompt = """
            Extract all text from this image perfectly.
            - Preserve the layout (paragraphs, tables) as much as possible.
            - If Thai text is present, ensure correct spelling.
            - Do NOT print ASCII borders like |---| if possible, use spacing.
            """
        
        response = model.generate_content([prompt, image])
        
        # ถ้าเป็น CSV ไม่ต้อง clean เส้นตาราง (เดี๋ยวข้อมูลหาย)
        if output_format == "csv":
            return response.text.replace("```csv", "").replace("```", "").strip()
        else:
            return clean_ocr_text(response.text)
            
    except Exception as e:
        return f"[Error on this page: {str(e)}]"

def create_word_docx(text_list):
    doc = Document()
    for i, text in enumerate(text_list):
        # ข้ามหน้าว่าง (ที่เป็น Excel)
        if not text or "Error" in text or "," in text and "\n" in text and len(text.split('\n')[0].split(',')) > 1:
             # Logic เช็ค CSV คร่าวๆ (อาจปรับปรุงได้) แต่เบื้องต้นเราเช็คจาก State ดีกว่า
             pass 
             
        doc.add_heading(f'Page {i+1}', level=1)
        doc.add_paragraph(text)
        doc.add_page_break()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_excel_from_results(results_map):
    """สร้าง Excel จาก Dictionary {page_index: csv_text}"""
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        has_data = False
        for page_idx, csv_text in results_map.items():
            if not csv_text: continue
            try:
                df = pd.read_csv(io.StringIO(csv_text))
                sheet_name = f"Page_{page_idx+1}"
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                has_data = True
            except:
                pass
        
        if not has_data:
            pd.DataFrame({"Info": ["No table data selected"]}).to_excel(writer, sheet_name="Info")
            
    buffer.seek(0)
    return buffer

def render_ocr_mode():
    # --- Session State ---
    if 'ocr_results' not in st.session_state: st.session_state['ocr_results'] = [] 
    if 'ocr_types' not in st.session_state: st.session_state['ocr_types'] = [] # เก็บประเภท (text/csv) ของแต่ละหน้า
    if 'ocr_images' not in st.session_state: st.session_state['ocr_images'] = []
    if 'current_page_index' not in st.session_state: st.session_state['current_page_index'] = 0
    if 'processed_file_id' not in st.session_state: st.session_state['processed_file_id'] = None

    # 1. แผงควบคุม (Expander)
    with st.expander("⚙️ แผงควบคุม (Control Panel)", expanded=True):
        col_key, col_model = st.columns([1, 1])
        with col_key:
            api_key = None
            if "GEMINI_API_KEY" in st.secrets:
                api_key = st.secrets["GEMINI_API_KEY"]
                st.success("✅ API Key Connected")
            else:
                api_key = st.text_input("🔑 Gemini API Key", type="password")
        with col_model:
            selected_model = None
            if api_key:
                model_options = get_available_models(api_key)
                if model_options:
                    default_idx = 0
                    for i, name in enumerate(model_options):
                        if "flash" in name and "exp" not in name:
                            default_idx = i; break
                    selected_model = st.selectbox("🤖 เลือก AI Model", model_options, index=default_idx)

        uploaded_file = st.file_uploader("📄 อัปโหลดไฟล์ PDF (AI OCR)", type=["pdf"])

        if uploaded_file and api_key and selected_model:
            # Check File Change
            if st.session_state['processed_file_id'] != uploaded_file.file_id:
                # Reset แต่ยังไม่เคลียร์ค่า จนกว่าจะกดเริ่ม
                pass

            # --- TABS ---
            tab_batch, tab_select = st.tabs(["🚀 แปลงทั้งหมด (Batch)", "👁️ เลือกเฉพาะหน้า (Selective)"])

            # TAB 1: BATCH (Text Only)
            with tab_batch:
                st.info("ℹ️ แปลงทุกหน้าเป็นข้อความ (Word)")
                if st.button("🚀 เริ่ม OCR ทุกหน้า", type="primary", use_container_width=True):
                    with st.spinner("📦 กำลังแยกหน้า PDF..."):
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        temp_images = []
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            pix = page.get_pixmap(dpi=150)
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            temp_images.append(img)
                        
                        st.session_state['ocr_images'] = temp_images
                        st.session_state['ocr_results'] = [""] * len(temp_images)
                        st.session_state['ocr_types'] = ["text"] * len(temp_images) # Default text
                        st.session_state['processed_file_id'] = uploaded_file.file_id
                        st.session_state['current_page_index'] = 0

                    progress_bar = st.progress(0, text="กำลังเริ่ม OCR...")
                    total_pages = len(st.session_state['ocr_images'])
                    
                    for i, img in enumerate(st.session_state['ocr_images']):
                        progress_bar.progress((i) / total_pages, text=f"🔍 กำลังอ่านหน้า {i+1}/{total_pages}...")
                        text_result = ocr_single_image(api_key, img, selected_model, "text")
                        st.session_state['ocr_results'][i] = text_result
                    
                    progress_bar.progress(1.0, text="เสร็จเรียบร้อย! (พับกล่องนี้เพื่อดูผลลัพธ์)")
                    st.rerun()

            # TAB 2: SELECTIVE (Text/Excel)
            with tab_select:
                st.info("ℹ️ เลือกหน้าและระบุว่าเป็นตาราง (Excel) หรือข้อความ (Word)")
                
                if 'ocr_preview_imgs' not in st.session_state or st.session_state.get('ocr_preview_fid') != uploaded_file.file_id:
                    with st.spinner("🖼️ สร้างภาพตัวอย่าง..."):
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        previews = []
                        for i in range(len(doc)):
                            page = doc.load_page(i)
                            pix = page.get_pixmap(dpi=72)
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            previews.append(img)
                        st.session_state['ocr_preview_imgs'] = previews
                        st.session_state['ocr_preview_fid'] = uploaded_file.file_id

                with st.form("ocr_select_form"):
                    images = st.session_state['ocr_preview_imgs']
                    cols = st.columns(4)
                    selection_map = {} # {index: type}
                    
                    for i, img in enumerate(images):
                        col = cols[i % 4]
                        with col:
                            st.image(img, use_container_width=True)
                            is_selected = st.checkbox(f"เลือกหน้า {i+1}", key=f"ocr_sel_{i}")
                            is_table = st.toggle(f"เป็นตาราง?", key=f"ocr_tbl_{i}")
                            
                            if is_selected:
                                selection_map[i] = "csv" if is_table else "text"
                    
                    st.markdown("---")
                    submitted = st.form_submit_button("✅ เริ่ม OCR เฉพาะหน้าที่เลือก", type="primary", use_container_width=True)

                if submitted:
                    if not selection_map:
                        st.warning("กรุณาเลือกอย่างน้อย 1 หน้า")
                    else:
                        # Reset
                        st.session_state['ocr_results'] = []
                        st.session_state['ocr_types'] = []
                        st.session_state['ocr_images'] = []
                        st.session_state['current_page_index'] = 0
                        st.session_state['processed_file_id'] = uploaded_file.file_id
                        
                        progress_bar = st.progress(0, text="เริ่มทำงาน...")
                        
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        total_sel = len(selection_map)
                        current_step = 0
                        
                        # Process in order
                        for page_idx, mode in sorted(selection_map.items()):
                            current_step += 1
                            progress_bar.progress((current_step / total_sel), text=f"🔍 กำลังอ่านหน้า {page_idx+1} ({mode})...")
                            
                            page = doc.load_page(page_idx)
                            pix = page.get_pixmap(dpi=150)
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            
                            st.session_state['ocr_images'].append(img)
                            st.session_state['ocr_types'].append(mode)
                            
                            text_res = ocr_single_image(api_key, img, selected_model, mode)
                            st.session_state['ocr_results'].append(text_res)
                        
                        progress_bar.progress(1.0, text="เสร็จเรียบร้อย! (พับกล่องนี้เพื่อดูผลลัพธ์)")
                        st.rerun()

    # 2. ส่วนแสดงผล (Outside Expander)
    if st.session_state.get('processed_file_id') == uploaded_file.file_id if uploaded_file else False:
        if st.session_state.get('ocr_results'):
            
            st.markdown("### 📄 ผลลัพธ์และดาวน์โหลด (Result & Export)")
            
            # --- Export Buttons ---
            # เช็คว่ามีผลลัพธ์แบบไหนบ้าง
            has_text = "text" in st.session_state['ocr_types']
            has_csv = "csv" in st.session_state['ocr_types']
            
            col_d1, col_d2 = st.columns(2)
            if has_text:
                with col_d1:
                    # Filter only text results
                    text_data = [res for res, type_ in zip(st.session_state['ocr_results'], st.session_state['ocr_types']) if type_ == "text"]
                    docx_file = create_word_docx(text_data)
                    st.download_button("💾 Export Word (.docx)", docx_file, "ocr_result.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
            
            if has_csv:
                with col_d2:
                    # Filter only csv results (map with page index)
                    csv_map = {i: res for i, (res, type_) in enumerate(zip(st.session_state['ocr_results'], st.session_state['ocr_types'])) if type_ == "csv"}
                    excel_file = create_excel_from_results(csv_map)
                    st.download_button("📊 Export Excel (.xlsx)", excel_file, "ocr_tables.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="secondary", use_container_width=True)

            st.markdown("---")

            # --- Synced View Controller ---
            total_pages = len(st.session_state['ocr_images'])
            col_prev, col_nav_info, col_next = st.columns([1, 4, 1])
            
            with col_prev:
                if st.button("⬅️ ก่อนหน้า", use_container_width=True, disabled=(st.session_state['current_page_index'] == 0)):
                    st.session_state['current_page_index'] -= 1
                    st.rerun()
            with col_nav_info:
                curr = st.session_state['current_page_index']
                type_label = "📝 ข้อความ" if st.session_state['ocr_types'][curr] == "text" else "📊 ตาราง (CSV)"
                st.markdown(f"<div style='text-align: center; padding-top: 5px; font-weight: bold;'>หน้า {curr + 1} / {total_pages} • {type_label}</div>", unsafe_allow_html=True)
            with col_next:
                if st.button("ถัดไป ➡️", use_container_width=True, disabled=(st.session_state['current_page_index'] == total_pages - 1)):
                    st.session_state['current_page_index'] += 1
                    st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)
            col_left_view, col_right_view = st.columns([1, 1])
            curr_idx = st.session_state['current_page_index']
            
            with col_left_view:
                st.info("👁️ ต้นฉบับ")
                if curr_idx < len(st.session_state['ocr_images']):
                    st.image(st.session_state['ocr_images'][curr_idx], use_container_width=True)

            with col_right_view:
                res_type = st.session_state['ocr_types'][curr_idx]
                if res_type == "text":
                    st.success("📝 ผลลัพธ์ (แก้ไขได้)")
                else:
                    st.warning("📊 ผลลัพธ์ CSV (สำหรับ Excel)")
                    
                if curr_idx < len(st.session_state['ocr_results']):
                    edited_text = st.text_area(
                        label="ocr_output",
                        value=st.session_state['ocr_results'][curr_idx],
                        height=800,
                        label_visibility="collapsed",
                        key=f"text_area_{curr_idx}"
                    )
                    st.session_state['ocr_results'][curr_idx] = edited_text
