import streamlit as st
import google.generativeai as genai
import fitz  # PyMuPDF
from PIL import Image
import io
from docx import Document
import re
import pandas as pd

def get_available_models(api_key):
    try:
        genai.configure(api_key=api_key)
        all_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        return all_models
    except:
        return []

def parse_ai_response(raw_text):
    """
    แยกเนื้อหา:
    1. ข้อความทั่วไป (Clean Text) -> สำหรับ Word
    2. ข้อมูลตาราง (CSV List) -> สำหรับ Excel
    """
    if not raw_text: 
        return "", []

    # Regex ค้นหาข้อความที่อยู่ระหว่าง [[TABLE]]...[[/TABLE]]
    # re.DOTALL เพื่อให้ . ครอบคลุมบรรทัดใหม่ด้วย
    table_pattern = re.compile(r'\[\[TABLE\]\](.*?)\[\[/TABLE\]\]', re.DOTALL)
    
    found_tables = []
    
    # ฟังก์ชันสำหรับแทนที่ตารางในข้อความหลักด้วย Marker
    def replace_with_marker(match):
        csv_content = match.group(1).strip()
        if csv_content:
            found_tables.append(csv_content)
            return "\n[--- ตรวจพบตาราง: ดูรายละเอียดในไฟล์ Excel ---]\n"
        return ""

    # 1. สร้าง Clean Text (เอาตารางออกแล้วแปะป้ายแทน)
    clean_text = table_pattern.sub(replace_with_marker, raw_text)
    
    # ล้างบรรทัดว่างส่วนเกิน
    clean_text = re.sub(r'\n{3,}', '\n\n', clean_text).strip()

    return clean_text, found_tables

def ocr_single_image(api_key, image, model_name):
    try:
        genai.configure(api_key=api_key)
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        model = genai.GenerativeModel(model_name, safety_settings=safety_settings)
        
        # --- PROMPT สูตรพิเศษ: สั่งให้แยกตารางด้วยแท็ก ---
        prompt = """
        Analyze this image and extract content.
        1. **Text**: Extract normal text with original layout.
        2. **Tables**: If you see any data table, DO NOT format it as Markdown. 
           Instead, convert it to CSV format and wrap it strictly within [[TABLE]] and [[/TABLE]] tags.
           Example:
           [[TABLE]]
           Column1,Column2
           Val1,Val2
           [[/TABLE]]
        3. **Thai Language**: Ensure high accuracy.
        """
        
        response = model.generate_content([prompt, image])
        
        # ส่งค่ากลับเป็น Raw Text ก่อน เดี๋ยวไปแยกข้างนอก
        return response.text
        
    except Exception as e:
        return f"[Error: {str(e)}]"

def create_word_docx(text_list):
    doc = Document()
    for i, text in enumerate(text_list):
        doc.add_heading(f'Page {i+1}', level=1)
        doc.add_paragraph(text)
        doc.add_page_break()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_excel_from_tables(all_pages_tables):
    """
    all_pages_tables: list ของ list (แต่ละหน้าอาจมีหลายตาราง)
    Format: [ [table1_p1, table2_p1], [table1_p2], ... ]
    """
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        has_data = False
        
        for page_idx, tables in enumerate(all_pages_tables):
            for table_idx, csv_data in enumerate(tables):
                try:
                    # แปลง CSV String เป็น DataFrame
                    df = pd.read_csv(io.StringIO(csv_data))
                    
                    # ตั้งชื่อ Sheet: P1_T1 (หน้า 1 ตาราง 1)
                    sheet_name = f"P{page_idx+1}_Table{table_idx+1}"
                    
                    # เขียนลง Excel
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                    has_data = True
                except:
                    pass
        
        if not has_data:
            pd.DataFrame({"Message": ["ไม่พบตารางในเอกสาร"]}).to_excel(writer, sheet_name="NoTables")
            
    buffer.seek(0)
    return buffer

def render_ocr_mode():
    # --- Session State ---
    if 'ocr_results_text' not in st.session_state: st.session_state['ocr_results_text'] = [] 
    if 'ocr_results_tables' not in st.session_state: st.session_state['ocr_results_tables'] = [] 
    if 'ocr_images' not in st.session_state: st.session_state['ocr_images'] = []
    if 'current_page_index' not in st.session_state: st.session_state['current_page_index'] = 0
    if 'processed_file_id' not in st.session_state: st.session_state['processed_file_id'] = None

    # 1. แผงควบคุม (Expander)
    with st.expander("⚙️ แผงควบคุม (Control Panel)", expanded=True):
        col_key, col_model = st.columns([1, 1])
        with col_key:
            api_key = None
            if "GEMINI_API_KEY" in st.secrets:
                api_key = st.secrets["GEMINI_API_KEY"]
                st.success("✅ API Key Connected")
            else:
                api_key = st.text_input("🔑 Gemini API Key", type="password")
        with col_model:
            selected_model = None
            if api_key:
                model_options = get_available_models(api_key)
                if model_options:
                    default_idx = 0
                    for i, name in enumerate(model_options):
                        if "flash" in name and "exp" not in name:
                            default_idx = i; break
                    selected_model = st.selectbox("🤖 เลือก AI Model", model_options, index=default_idx)

        uploaded_file = st.file_uploader("📄 อัปโหลดไฟล์ PDF (AI OCR)", type=["pdf"])

        if uploaded_file and api_key and selected_model:
            # Check File Change
            if st.session_state['processed_file_id'] != uploaded_file.file_id:
                # Reset
                pass

            # --- TABS ---
            tab_batch, tab_select = st.tabs(["🚀 แปลงทั้งหมด (Batch)", "👁️ เลือกเฉพาะหน้า (Selective)"])

            # TAB 1: BATCH
            with tab_batch:
                st.info("ℹ️ อ่านทุกหน้า + แยกตารางให้อัตโนมัติ")
                if st.button("🚀 เริ่ม OCR ทุกหน้า", type="primary", use_container_width=True):
                    with st.spinner("📦 กำลังแยกหน้า PDF..."):
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        temp_images = []
                        for page_num in range(len(doc)):
                            page = doc.load_page(page_num)
                            pix = page.get_pixmap(dpi=150)
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            temp_images.append(img)
                        
                        st.session_state['ocr_images'] = temp_images
                        st.session_state['ocr_results_text'] = [""] * len(temp_images)
                        st.session_state['ocr_results_tables'] = [[]] * len(temp_images)
                        st.session_state['processed_file_id'] = uploaded_file.file_id
                        st.session_state['current_page_index'] = 0

                    progress_bar = st.progress(0, text="กำลังเริ่ม OCR...")
                    total_pages = len(st.session_state['ocr_images'])
                    
                    for i, img in enumerate(st.session_state['ocr_images']):
                        progress_bar.progress((i) / total_pages, text=f"🔍 กำลังอ่านหน้า {i+1}/{total_pages}...")
                        
                        # Call AI
                        raw_response = ocr_single_image(api_key, img, selected_model)
                        
                        # Parse: แยก Text กับ Tables
                        clean_text, tables = parse_ai_response(raw_response)
                        
                        st.session_state['ocr_results_text'][i] = clean_text
                        st.session_state['ocr_results_tables'][i] = tables
                    
                    progress_bar.progress(1.0, text="เสร็จเรียบร้อย! (พับกล่องนี้เพื่อดูผลลัพธ์)")
                    st.rerun()

            # TAB 2: SELECTIVE
            with tab_select:
                st.info("ℹ️ เลือกเฉพาะหน้า (ระบบจะแยกตารางให้อัตโนมัติเช่นกัน)")
                
                if 'ocr_preview_imgs' not in st.session_state or st.session_state.get('ocr_preview_fid') != uploaded_file.file_id:
                    with st.spinner("🖼️ สร้างภาพตัวอย่าง..."):
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        previews = []
                        for i in range(len(doc)):
                            page = doc.load_page(i)
                            pix = page.get_pixmap(dpi=72)
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            previews.append(img)
                        st.session_state['ocr_preview_imgs'] = previews
                        st.session_state['ocr_preview_fid'] = uploaded_file.file_id

                with st.form("ocr_select_form"):
                    images = st.session_state['ocr_preview_imgs']
                    cols = st.columns(4)
                    selected_indices = []
                    
                    for i, img in enumerate(images):
                        col = cols[i % 4]
                        with col:
                            st.image(img, use_container_width=True)
                            if st.checkbox(f"หน้า {i+1}", key=f"ocr_sel_{i}"):
                                selected_indices.append(i)
                    
                    st.markdown("---")
                    submitted = st.form_submit_button("✅ เริ่ม OCR เฉพาะหน้าที่เลือก", type="primary", use_container_width=True)

                if submitted:
                    if not selected_indices:
                        st.warning("กรุณาเลือกอย่างน้อย 1 หน้า")
                    else:
                        st.session_state['ocr_results_text'] = []
                        st.session_state['ocr_results_tables'] = []
                        st.session_state['ocr_images'] = []
                        st.session_state['current_page_index'] = 0
                        st.session_state['processed_file_id'] = uploaded_file.file_id
                        
                        progress_bar = st.progress(0, text="เริ่มทำงาน...")
                        
                        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                        total_sel = len(selected_indices)
                        current_step = 0
                        
                        selected_indices.sort()
                        
                        for idx, page_num in enumerate(selected_indices):
                            current_step += 1
                            progress_bar.progress((current_step / total_sel), text=f"🔍 กำลังอ่านหน้า {page_num+1} ({current_step}/{total_sel})...")
                            
                            page = doc.load_page(page_num)
                            pix = page.get_pixmap(dpi=150)
                            img = Image.open(io.BytesIO(pix.tobytes()))
                            
                            st.session_state['ocr_images'].append(img)
                            
                            # Call AI
                            raw_response = ocr_single_image(api_key, img, selected_model)
                            # Parse
                            clean_text, tables = parse_ai_response(raw_response)
                            
                            st.session_state['ocr_results_text'].append(clean_text)
                            st.session_state['ocr_results_tables'].append(tables)
                        
                        progress_bar.progress(1.0, text="เสร็จเรียบร้อย! (พับกล่องนี้เพื่อดูผลลัพธ์)")
                        st.rerun()

    # 2. ส่วนแสดงผล (Outside Expander)
    if st.session_state.get('processed_file_id') == uploaded_file.file_id if uploaded_file else False:
        if st.session_state.get('ocr_results_text'):
            
            st.markdown("### 📄 ผลลัพธ์ (Result & Export)")
            
            # --- Check Data ---
            has_text = any(st.session_state['ocr_results_text'])
            # เช็คว่ามีตารางอย่างน้อย 1 หน้าไหม
            has_tables = any(len(t) > 0 for t in st.session_state['ocr_results_tables'])
            
            # --- Export Buttons ---
            col_d1, col_d2 = st.columns(2)
            
            with col_d1:
                if has_text:
                    docx_file = create_word_docx(st.session_state['ocr_results_text'])
                    st.download_button("💾 Export Word (.docx)", docx_file, "ocr_result.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
            
            with col_d2:
                if has_tables:
                    excel_file = create_excel_from_tables(st.session_state['ocr_results_tables'])
                    st.download_button("📊 Export Tables (.xlsx)", excel_file, "ocr_tables.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="secondary", use_container_width=True)
                else:
                    st.info("ℹ️ ไม่พบตารางในเอกสาร (ปุ่มโหลด Excel จึงไม่แสดง)")

            st.markdown("---")

            # --- Synced View Controller ---
            total_pages = len(st.session_state['ocr_images'])
            col_prev, col_nav_info, col_next = st.columns([1, 4, 1])
            
            with col_prev:
                if st.button("⬅️ ก่อนหน้า", use_container_width=True, disabled=(st.session_state['current_page_index'] == 0)):
                    st.session_state['current_page_index'] -= 1
                    st.rerun()
            with col_nav_info:
                curr = st.session_state['current_page_index']
                # บอก User ว่าหน้านี้มีตารางไหม
                table_count = len(st.session_state['ocr_results_tables'][curr])
                status_msg = f"หน้า {curr + 1} / {total_pages}"
                if table_count > 0:
                    status_msg += f" (พบ {table_count} ตาราง ✅)"
                
                st.markdown(f"<div style='text-align: center; padding-top: 5px; font-weight: bold;'>{status_msg}</div>", unsafe_allow_html=True)
            with col_next:
                if st.button("ถัดไป ➡️", use_container_width=True, disabled=(st.session_state['current_page_index'] == total_pages - 1)):
                    st.session_state['current_page_index'] += 1
                    st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)
            col_left_view, col_right_view = st.columns([1, 1])
            curr_idx = st.session_state['current_page_index']
            
            with col_left_view:
                st.info("👁️ ต้นฉบับ")
                if curr_idx < len(st.session_state['ocr_images']):
                    st.image(st.session_state['ocr_images'][curr_idx], use_container_width=True)

            with col_right_view:
                st.success("📝 ข้อความหลัก (Main Text)")
                if curr_idx < len(st.session_state['ocr_results_text']):
                    edited_text = st.text_area(
                        label="ocr_output",
                        value=st.session_state['ocr_results_text'][curr_idx],
                        height=800,
                        label_visibility="collapsed",
                        key=f"text_area_{curr_idx}"
                    )
                    st.session_state['ocr_results_text'][curr_idx] = edited_text
